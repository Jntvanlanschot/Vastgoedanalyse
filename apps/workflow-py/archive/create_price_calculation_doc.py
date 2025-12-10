#!/usr/bin/env python3
"""
Create Word document with price calculation formulas and weights.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx niet geïnstalleerd. Installeer met: pip install python-docx")
    sys.exit(1)

def add_formula_paragraph(doc, text, bold=False, indent=0):
    """Add a paragraph with formula text."""
    p = doc.add_paragraph()
    p.style = 'Normal'
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    
    return p

def create_price_calculation_doc():
    """Create Word document with price calculation formulas."""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Overzicht Prijsberekening Vastgoedanalyse Tool', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Geoptimaliseerde Parameters
    doc.add_heading('Geoptimaliseerde Parameters', 1)
    add_formula_paragraph(doc, 'Resultaten van Bayesian Optimization (300 trials):', bold=True)
    add_formula_paragraph(doc, '• Within 10% accuracy: 54.82%')
    add_formula_paragraph(doc, '• MAPE (Mean Absolute Percentage Error): 11.52%')
    
    doc.add_paragraph()
    
    # 1. SIMILARITY SCORE BEREKENING
    doc.add_heading('1. Similarity Score Berekening', 1)
    add_formula_paragraph(doc, 'De similarity score bepaalt hoe goed een vergelijkbare woning matcht met de referentie woning. De score wordt berekend op basis van meerdere factoren met geoptimaliseerde weights.')
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Similarity Weights (Actief Gebruikt):', bold=True)
    
    # Weights table
    weights_data = [
        ['Factor', 'Weight', 'Percentage', 'Beschrijving'],
        ['Energielabel', '0.35', '35%', 'Meest belangrijke factor - bepaalt energie-efficiëntie'],
        ['Oppervlakte (m²)', '0.33', '33%', 'Verschil in woonoppervlakte'],
        ['Afstand/Locatie', '0.18', '18%', 'Geografische afstand (0-2km linear decay)'],
        ['Balkon/Terras', '0.11', '11%', 'Match op balkon of terras aanwezigheid'],
        ['Verkoopdatum', '0.11', '11%', 'Recency en tijd nabijheid van verkoopdatum'],
        ['Straat naam', '0.10', '10%', 'Exacte match of string similarity'],
        ['OSM Straat', '0.10', '10%', 'OpenStreetMap straat similarity'],
        ['Kamers', '0.05', '5%', 'Aantal kamers verschil'],
        ['Tuin', '0.02', '2%', 'Tuin aanwezigheid match'],
        ['Bouwjaar', '0.01', '1%', 'Verschil in bouwjaar'],
        ['Gracht Penalty', '0.0035', '0.35%', 'Penalty multiplier als één gracht is en andere niet'],
    ]
    
    table = doc.add_table(rows=len(weights_data), cols=4)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(weights_data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            if i == 0:  # Header row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Totaal Base Similarity: 10% + 10% + 33% + 18% + 2% + 5% + 11% + 11% + 1% = 101%', bold=True)
    
    doc.add_paragraph()
    doc.add_heading('Berekeningsmethode', 2)
    
    add_formula_paragraph(doc, '1. Base Similarity (zonder energielabel):', bold=True)
    add_formula_paragraph(doc, '   Alle factoren behalve energielabel worden opgeteld', indent=0.25)
    add_formula_paragraph(doc, '   Maximaal: 10% + 10% + 33% + 18% + 2% + 5% + 11% + 11% + 1% = 101%', indent=0.25)
    add_formula_paragraph(doc, '   Score wordt genormaliseerd naar 0-1 door te delen door maximum score', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, '2. Verkoopdatum Similarity:', bold=True)
    add_formula_paragraph(doc, '   Formule: proximity_score = exp(-days_diff / 180)', indent=0.25)
    add_formula_paragraph(doc, '   • Zelfde datum: score = 1.0', indent=0.25)
    add_formula_paragraph(doc, '   • Verschil in dagen wordt gebruikt met exponentiële decay', indent=0.25)
    add_formula_paragraph(doc, '   • Decay factor = 180 dagen (50% score bij ~125 dagen verschil)', indent=0.25)
    add_formula_paragraph(doc, '   • Recency bonus: +0.05 als beide verkoopdata recent zijn (< 1 jaar)', indent=0.25)
    add_formula_paragraph(doc, '   • Maximum verschil: 730 dagen (2 jaar), daarboven score = 0', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, '3. Bouwjaar Similarity:', bold=True)
    add_formula_paragraph(doc, '   Formule: score = exp(-year_diff / 20)', indent=0.25)
    add_formula_paragraph(doc, '   • Zelfde jaar: score = 1.0', indent=0.25)
    add_formula_paragraph(doc, '   • Verschil in jaren wordt gebruikt met exponentiële decay', indent=0.25)
    add_formula_paragraph(doc, '   • Decay factor = 20 jaar (50% score bij ~14 jaar verschil)', indent=0.25)
    add_formula_paragraph(doc, '   • Bij ontbrekende data: score = 0.5 (neutraal)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, '4. Energielabel Similarity:', bold=True)
    add_formula_paragraph(doc, '   Gebruikt exponentiële decay functie op basis van label verschil', indent=0.25)
    add_formula_paragraph(doc, '   • Zelfde label: score = 1.0', indent=0.25)
    add_formula_paragraph(doc, '   • Verschil van 1 label: score = 0.9', indent=0.25)
    add_formula_paragraph(doc, '   • Verschil van 2 labels: score = 0.75', indent=0.25)
    add_formula_paragraph(doc, '   • Verschil van 3 labels: score = 0.6', indent=0.25)
    add_formula_paragraph(doc, '   • Grotere verschillen: exponentiële decay met minimum 0.4', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, '5. Gecombineerde Score:', bold=True)
    add_formula_paragraph(doc, '   Final Score = (35% × Energy Label Similarity) + (65% × Base Similarity)', indent=0.25)
    add_formula_paragraph(doc, '   Waarbij Base Similarity = genormaliseerde som van alle andere factoren', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, '6. Gracht Penalty:', bold=True)
    add_formula_paragraph(doc, '   Als referentie woning een gracht is maar vergelijkbare niet (of vice versa):', indent=0.25)
    add_formula_paragraph(doc, '   Final Score = Score × 0.0035 (zware penalty)', indent=0.25)
    add_formula_paragraph(doc, '   Dit voorkomt dat grachten met niet-grachten worden vergeleken', indent=0.25)
    
    doc.add_page_break()
    
    # 2. PRIJSBEREKENING PARAMETERS
    doc.add_heading('2. Prijsberekening Parameters', 1)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Selectie Criteria:', bold=True)
    
    params_data = [
        ['Parameter', 'Waarde', 'Beschrijving'],
        ['TOP_N', '12', 'Aantal top vergelijkbare woningen gebruikt voor prijsberekening'],
        ['MIN_SCORE', '0.55', 'Minimum similarity score (55%) - woningen onder deze score worden uitgesloten'],
        ['USE_SCORE_SQUARED', 'True', 'Gebruik score² als weight (geeft meer gewicht aan hoge scores)'],
    ]
    
    table2 = doc.add_table(rows=len(params_data), cols=3)
    table2.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(params_data):
        for j, cell_data in enumerate(row_data):
            cell = table2.rows[i].cells[j]
            cell.text = cell_data
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_heading('Prijsberekening Stappen', 2)
    
    add_formula_paragraph(doc, 'Stap 1: Filtering', bold=True)
    add_formula_paragraph(doc, '   • Neem top 15 matches uit de dataset', indent=0.25)
    add_formula_paragraph(doc, '   • Filter op minimum score ≥ 0.55', indent=0.25)
    add_formula_paragraph(doc, '   • Neem top 12 woningen die voldoen aan minimum score', indent=0.25)
    add_formula_paragraph(doc, '   • Fallback: Als geen woningen voldoen, gebruik top 5 zonder score filter', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Stap 2: Energy Label Correctie', bold=True)
    add_formula_paragraph(doc, '   Formule: Corrected Price = Original Price × (1 + (ref_factor - comp_factor))', indent=0.25)
    add_formula_paragraph(doc, '   Waarbij:', indent=0.25)
    add_formula_paragraph(doc, '   • ref_factor = energy label adjustment van referentie woning', indent=0.5)
    add_formula_paragraph(doc, '   • comp_factor = energy label adjustment van vergelijkbare woning', indent=0.5)
    add_formula_paragraph(doc, '   Energy label adjustments:', indent=0.25)
    add_formula_paragraph(doc, '   • Label A: +8%', indent=0.5)
    add_formula_paragraph(doc, '   • Label B: +4%', indent=0.5)
    add_formula_paragraph(doc, '   • Label C: 0% (baseline)', indent=0.5)
    add_formula_paragraph(doc, '   • Label D: -5%', indent=0.5)
    add_formula_paragraph(doc, '   • Label E: -8%', indent=0.5)
    add_formula_paragraph(doc, '   • Label F: -12%', indent=0.5)
    add_formula_paragraph(doc, '   • Label G: -15%', indent=0.5)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Stap 3: Prijs per m² Berekening', bold=True)
    add_formula_paragraph(doc, '   Formule: Price per m² = Corrected Price / Area (m²)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Stap 4: Weighted Average', bold=True)
    add_formula_paragraph(doc, '   Elke prijs per m² krijgt een weight gebaseerd op similarity score:', indent=0.25)
    add_formula_paragraph(doc, '   • Weight = Score² (als USE_SCORE_SQUARED = True)', indent=0.5)
    add_formula_paragraph(doc, '   • Weight = Score (als USE_SCORE_SQUARED = False)', indent=0.5)
    add_formula_paragraph(doc, '   Formule: Weighted Average = Σ(Price per m² × Weight) / Σ(Weight)', indent=0.25)
    add_formula_paragraph(doc, '   Waarbij Σ = som over alle geselecteerde vergelijkbare woningen', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Stap 5: Finale Prijs (Neutraal Scenario)', bold=True)
    add_formula_paragraph(doc, '   Formule: Final Price = Weighted Average Price per m² × Reference Area (m²)', indent=0.25)
    
    doc.add_page_break()
    
    # 3. DRIE SCENARIO'S
    doc.add_heading('3. Drie Scenario\'s (Bounds)', 1)
    add_formula_paragraph(doc, 'Gebaseerd op financiële literatuur worden drie scenario\'s berekend:')
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Conservatief Scenario (P25):', bold=True)
    add_formula_paragraph(doc, '   Methode: 25e percentiel van alle prijzen per m²', indent=0.25)
    add_formula_paragraph(doc, '   Betekenis: 25% van de vergelijkbare woningen heeft een lagere prijs', indent=0.25)
    add_formula_paragraph(doc, '   Formule: Conservatief Prijs = P25(prijzen per m²) × Reference Area (m²)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Neutraal Scenario:', bold=True)
    add_formula_paragraph(doc, '   Methode: Gewogen gemiddelde (zoals hierboven beschreven)', indent=0.25)
    add_formula_paragraph(doc, '   Betekenis: Meest waarschijnlijke prijs gebaseerd op alle factoren', indent=0.25)
    add_formula_paragraph(doc, '   Formule: Neutraal Prijs = Weighted Average Price per m² × Reference Area (m²)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Optimistisch Scenario (P75):', bold=True)
    add_formula_paragraph(doc, '   Methode: 75e percentiel van alle prijzen per m²', indent=0.25)
    add_formula_paragraph(doc, '   Betekenis: 75% van de vergelijkbare woningen heeft een lagere prijs', indent=0.25)
    add_formula_paragraph(doc, '   Formule: Optimistisch Prijs = P75(prijzen per m²) × Reference Area (m²)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Confidence Interval:', bold=True)
    add_formula_paragraph(doc, '   • 50% Confidence Interval: P25 tot P75', indent=0.25)
    add_formula_paragraph(doc, '   • Betekenis: 50% kans dat de werkelijke prijs binnen dit bereik valt', indent=0.25)
    add_formula_paragraph(doc, '   • Conform: Financiële standaarden voor property valuations', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Fallback Logica:', bold=True)
    add_formula_paragraph(doc, '   • ≥3 vergelijkbare woningen: Gebruik percentielen (P25/P75)', indent=0.25)
    add_formula_paragraph(doc, '   • 2 woningen: Min/max met 10% marge', indent=0.25)
    add_formula_paragraph(doc, '   • 1 woning: ±12% range (standaard in property valuations)', indent=0.25)
    
    doc.add_page_break()
    
    # 4. SAMENVATTING
    doc.add_heading('4. Samenvatting', 1)
    
    add_formula_paragraph(doc, 'Similarity Score:', bold=True)
    add_formula_paragraph(doc, '   • 35% energielabel + 65% andere factoren', indent=0.25)
    add_formula_paragraph(doc, '   • Andere factoren: straat (10%), OSM straat (10%), oppervlakte (33%), afstand (18%),', indent=0.25)
    add_formula_paragraph(doc, '     balkon (11%), verkoopdatum (11%), kamers (5%), tuin (2%), bouwjaar (1%)', indent=0.25)
    add_formula_paragraph(doc, '   • Maximaal 100% (1.0)', indent=0.25)
    add_formula_paragraph(doc, '   • Gracht penalty van 0.35% bij mismatch', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Prijsberekening:', bold=True)
    add_formula_paragraph(doc, '   • Top 12 woningen met score ≥ 55%', indent=0.25)
    add_formula_paragraph(doc, '   • Gewogen gemiddelde met score² als weight', indent=0.25)
    add_formula_paragraph(doc, '   • Energy label correctie toegepast op elke prijs', indent=0.25)
    add_formula_paragraph(doc, '   • Drie scenario\'s: Conservatief (P25), Neutraal (gewogen gemiddelde), Optimistisch (P75)', indent=0.25)
    
    doc.add_paragraph()
    add_formula_paragraph(doc, 'Output:', bold=True)
    add_formula_paragraph(doc, '   • Conservatief: 25e percentiel', indent=0.25)
    add_formula_paragraph(doc, '   • Neutraal: Gewogen gemiddelde', indent=0.25)
    add_formula_paragraph(doc, '   • Optimistisch: 75e percentiel', indent=0.25)
    add_formula_paragraph(doc, '   • 50% confidence interval tussen P25 en P75', indent=0.25)
    
    # Save document
    output_path = Path('PRIJSBEREKENING_FORMULES.docx')
    doc.save(str(output_path))
    print(f"Word document opgeslagen: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_price_calculation_doc()
    except ImportError as e:
        print(f"Fout: {e}")
        print("Installeer python-docx met: pip install python-docx")
        sys.exit(1)

